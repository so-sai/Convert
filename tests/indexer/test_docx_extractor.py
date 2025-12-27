"""
DOCX Extractor Tests - TDD GREEN Phase
Task 6.4 Extraction Pipeline

NOTE: Using Python fallback (python-docx) while Rust module builds.
Rust implementation (docx_rs) will be swapped in once build completes.
"""

import pytest
from pathlib import Path
import tempfile
import zipfile
import xml.etree.ElementTree as ET
import time

# Try to import python-docx for test data creation
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


@pytest.fixture
def sample_docx(tmp_path):
    """Create a minimal valid DOCX file for testing"""
    if not PYTHON_DOCX_AVAILABLE:
        pytest.skip("python-docx not available")
    
    docx_path = tmp_path / "test.docx"
    
    # Create DOCX with python-docx
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document.")
    doc.save(str(docx_path))
    
    return docx_path


@pytest.fixture
def complex_docx(tmp_path):
    """Create DOCX with tables and formatting"""
    if not PYTHON_DOCX_AVAILABLE:
        pytest.skip("python-docx not available")
    
    docx_path = tmp_path / "complex.docx"
    
    doc = Document()
    doc.add_paragraph("Document with table")
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Data 1"
    table.cell(1, 1).text = "Data 2"
    
    doc.add_paragraph("After table")
    
    doc.save(str(docx_path))
    
    return docx_path


@pytest.fixture
def corrupted_docx(tmp_path):
    """Create a corrupted DOCX file"""
    docx_path = tmp_path / "corrupted.docx"
    # Write invalid data
    docx_path.write_bytes(b"PK\x03\x04corrupted data")
    return docx_path


@pytest.mark.skipif(not PYTHON_DOCX_AVAILABLE, reason="python-docx not installed")
def test_T25_01_docx_extraction_basic(sample_docx):
    """T25.01: Extract text from valid DOCX file"""
    from src.core.indexer.extractors.docx_extractor import extract_docx
    
    result = extract_docx(sample_docx)
    
    assert len(result.segments) >= 2, f"Expected at least 2 segments, got {len(result.segments)}"
    assert result.processing_time_ms > 0
    assert len(result.errors) == 0
    assert not result.truncated
    assert result.success
    
    # Check text content
    all_text = " ".join(seg.text for seg in result.segments)
    assert "Hello World" in all_text
    assert "test document" in all_text
    
    print(f"\n   ✅ T25.01: Extracted {len(result.segments)} segments in {result.processing_time_ms:.2f}ms")


@pytest.mark.skipif(not PYTHON_DOCX_AVAILABLE, reason="python-docx not installed")
def test_T25_02_docx_extraction_complex(complex_docx):
    """T25.02: Handle complex formatting (tables, lists)"""
    from src.core.indexer.extractors.docx_extractor import extract_docx
    
    result = extract_docx(complex_docx)
    
    assert result.success
    assert len(result.segments) > 0
    
    # Should have extracted table content
    all_text = " ".join(seg.text for seg in result.segments)
    assert "Header 1" in all_text or "Data 1" in all_text
    
    # Check metadata
    assert "paragraph_count" in result.metadata
    assert "table_count" in result.metadata
    assert result.metadata["table_count"] >= 1
    
    print(f"\n   ✅ T25.02: Complex DOCX handled: {result.metadata}")


@pytest.mark.skipif(not PYTHON_DOCX_AVAILABLE, reason="python-docx not installed")
def test_T25_03_docx_extraction_invalid_file(corrupted_docx):
    """T25.03: Handle invalid DOCX gracefully"""
    from src.core.indexer.extractors.docx_extractor import extract_docx
    
    result = extract_docx(corrupted_docx)
    
    assert len(result.errors) > 0
    assert not result.success
    assert any("corrupt" in err.message.lower() or "failed" in err.message.lower() 
               for err in result.errors)
    
    print(f"\n   ✅ T25.03: Corrupted DOCX handled: {result.errors[0].message}")


@pytest.mark.skipif(not PYTHON_DOCX_AVAILABLE, reason="python-docx not installed")
def test_T25_04_docx_extraction_performance(sample_docx):
    """T25.04: Verify extraction performance < 50ms P95"""
    from src.core.indexer.extractors.docx_extractor import extract_docx
    
    times = []
    for _ in range(100):
        start = time.perf_counter()
        result = extract_docx(sample_docx)
        elapsed_ms = (time.perf_counter() - start) * 1000
        times.append(elapsed_ms)
        
        assert result.success
    
    times.sort()
    p50 = times[49]
    p95 = times[94]
    
