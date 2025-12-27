"""
DOCX_EXTRACTOR.PY - Python Fallback DOCX Text Extraction
Task 6.4 - Sprint 6 Background Services

Fallback implementation using python-docx while Rust module builds.
This provides immediate functionality while maintaining the same interface.

NOTE: Rust implementation (docx_rs) is preferred for production due to:
- 3-5x better performance
- No-GIL compatibility
- Lower memory footprint

This fallback ensures tests can run immediately.
"""

import time
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from .result import ExtractionResult, TextSegment, ExtractionError


class DOCXExtractor:
    """
    Python-docx based DOCX text extractor (fallback).
    
    Features:
    - Text extraction with paragraph tracking
    - Table content extraction
    - Graceful degradation for corrupted DOCX
    - Performance tracking
    
    Note:
        This is a fallback implementation. The Rust implementation
        (docx_rs) provides better performance and No-GIL compatibility.
    """
    
    VERSION = "1.0.0-python-fallback"
    EXTRACTOR_NAME = "docx_python"
    
    def __init__(self):
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx not available. Install with: pip install python-docx"
            )
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """
        Extract text and metadata from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ExtractionResult with text segments and metadata
        """
        start_time = time.perf_counter()
        result = ExtractionResult(
            extractor=self.EXTRACTOR_NAME,
            version=self.VERSION
        )
        
        # Get file size
        try:
            result.file_size_bytes = file_path.stat().st_size
        except Exception:
            result.file_size_bytes = 0
        
        # Open DOCX
        try:
            doc = Document(str(file_path))
        except FileNotFoundError:
            result.add_error(
                ExtractionError.FILE_NOT_FOUND,
                f"DOCX file not found: {file_path}",
                recoverable=False
            )
            result.processing_time_ms = (time.perf_counter() - start_time) * 1000
            return result
        except Exception as e:
            result.add_error(
                ExtractionError.CORRUPTED,
                f"Failed to open DOCX: {e}",
                recoverable=False
            )
            result.processing_time_ms = (time.perf_counter() - start_time) * 1000
            return result
        
        try:
            # Extract paragraphs
            para_count = 0
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:  # Only add non-empty paragraphs
                    result.segments.append(TextSegment(
                        text=text,
                        page=None,
                        section=f"paragraph_{idx}",
                        confidence=1.0
                    ))
                    para_count += 1
            
            # Extract tables
            table_count = 0
            for table_idx, table in enumerate(doc.tables):
                table_text_parts = []
                for row in table.rows:
                    row_text = "\t".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text_parts.append(row_text)
                
                if table_text_parts:
                    table_text = "\n".join(table_text_parts)
                    result.segments.append(TextSegment(
                        text=table_text,
                        page=None,
                        section=f"table_{table_idx}",
                        confidence=1.0
                    ))
                    table_count += 1
            
            # Add metadata
            result.metadata = {
                "paragraph_count": para_count,
                "table_count": table_count,
            }
            
            # Try to extract core properties
            try:
                core_props = doc.core_properties
                if core_props.title:
                    result.metadata["title"] = core_props.title
                if core_props.author:
                    result.metadata["author"] = core_props.author
                if core_props.subject:
                    result.metadata["subject"] = core_props.subject
                if core_props.created:
                    result.metadata["created"] = str(core_props.created)
                if core_props.modified:
                    result.metadata["modified"] = str(core_props.modified)
            except Exception:
                # Metadata extraction is non-critical
                pass
                
        except Exception as e:
            result.add_error(
                ExtractionError.CORRUPTED,
                f"Extraction failed: {e}",
                recoverable=len(result.segments) > 0
            )
        finally:
            result.processing_time_ms = (time.perf_counter() - start_time) * 1000
        
        return result


def extract_docx(file_path: Path) -> ExtractionResult:
    """
    Convenience function for DOCX extraction.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        ExtractionResult
    """
    extractor = DOCXExtractor()
    return extractor.extract(file_path)
